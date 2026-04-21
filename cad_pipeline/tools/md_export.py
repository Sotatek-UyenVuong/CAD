"""md_export.py — Markdown export helpers (PDF/DOCX via Pandoc)."""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path


def inject_front_matter(
    markdown: str,
    title: str = "Report",
    author: str = "CAD Pipeline",
    date: str | None = None,
    toc: bool = False,
) -> str:
    lines = [
        "---",
        f'title: "{title}"',
        f'author: "{author}"',
    ]
    if date:
        lines.append(f'date: "{date}"')
    if toc:
        lines += ["toc: true", "toc-depth: 2"]
    lines += ["---", ""]
    return "\n".join(lines) + (markdown or "")


def convert_to_pdf(
    md_path: str | Path,
    out_path: str | Path,
    title: str = "Report",
    author: str = "CAD Pipeline",
    date: str | None = None,
    toc: bool = False,
) -> bool:
    md_path = Path(md_path)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if not md_path.exists():
        return False
    if shutil.which("pandoc") is None:
        return False

    src_text = md_path.read_text(encoding="utf-8")
    enriched = inject_front_matter(src_text, title=title, author=author, date=date, toc=toc)
    tmp_md = md_path.with_suffix(".frontmatter.md")
    tmp_md.write_text(enriched, encoding="utf-8")

    cmd = [
        "pandoc",
        str(tmp_md),
        "-o",
        str(out_path),
        "--from",
        "markdown+pipe_tables+grid_tables+fenced_code_blocks+smart",
        "--pdf-engine=xelatex",
        "-V",
        "geometry:margin=1in",
        "-V",
        "mainfont=DejaVu Serif",
        "-V",
        "sansfont=DejaVu Sans",
        "-V",
        "monofont=DejaVu Sans Mono",
        "-V",
        "fontsize=11pt",
        "-V",
        "linestretch=1.2",
    ]
    try:
        proc = subprocess.run(cmd, check=False, capture_output=True, text=True)
        return proc.returncode == 0 and out_path.exists()
    finally:
        tmp_md.unlink(missing_ok=True)


def convert_to_docx(
    md_path: str | Path,
    out_path: str | Path,
    title: str = "Report",
    author: str = "CAD Pipeline",
    date: str | None = None,
    toc: bool = False,
    reference_doc: str | Path | None = None,
) -> bool:
    md_path = Path(md_path)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if not md_path.exists():
        return False
    if shutil.which("pandoc") is None:
        return _convert_to_docx_python(md_path, out_path, title=title)

    src_text = md_path.read_text(encoding="utf-8")
    enriched = inject_front_matter(src_text, title=title, author=author, date=date, toc=toc)
    tmp_md = md_path.with_suffix(".frontmatter.md")
    tmp_md.write_text(enriched, encoding="utf-8")

    cmd = [
        "pandoc",
        str(tmp_md),
        "-o",
        str(out_path),
    ]
    if reference_doc:
        ref_path = Path(reference_doc)
        if ref_path.exists():
            cmd.extend(["--reference-doc", str(ref_path)])
    try:
        proc = subprocess.run(cmd, check=False, capture_output=True, text=True)
        return proc.returncode == 0 and out_path.exists()
    finally:
        tmp_md.unlink(missing_ok=True)


def _convert_to_docx_python(
    md_path: Path,
    out_path: Path,
    title: str,
) -> bool:
    """Fallback DOCX export using python-docx when pandoc is unavailable."""
    try:
        from docx import Document  # type: ignore
    except Exception:
        return False

    try:
        text = md_path.read_text(encoding="utf-8")
    except Exception:
        return False

    def _strip_inline_md(s: str) -> str:
        v = s.strip()
        # links: [text](url) -> text
        v = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", v)
        # inline code
        v = re.sub(r"`([^`]+)`", r"\1", v)
        # bold / italic markers
        v = re.sub(r"\*\*([^*]+)\*\*", r"\1", v)
        v = re.sub(r"__([^_]+)__", r"\1", v)
        v = re.sub(r"\*([^*]+)\*", r"\1", v)
        v = re.sub(r"_([^_]+)_", r"\1", v)
        return v

    def _split_table_row(line: str) -> list[str]:
        row = line.strip().strip("|")
        return [_strip_inline_md(c) for c in row.split("|")]

    def _is_separator_row(line: str) -> bool:
        row = line.strip().strip("|").replace(" ", "")
        if not row:
            return False
        return all(ch in "-:|" for ch in row)

    lines = [ln.rstrip("\n") for ln in text.splitlines()]
    doc = Document()
    doc.add_heading(title or "Report", level=1)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # Table block
        if stripped.startswith("|") and stripped.endswith("|"):
            block: list[str] = []
            while i < len(lines):
                cur = lines[i].strip()
                if cur.startswith("|") and cur.endswith("|"):
                    block.append(cur)
                    i += 1
                    continue
                break
            data_rows = [r for r in block if not _is_separator_row(r)]
            if not data_rows:
                continue
            parsed = [_split_table_row(r) for r in data_rows]
            col_count = max((len(r) for r in parsed), default=0)
            if col_count == 0:
                continue
            table = doc.add_table(rows=0, cols=col_count)
            try:
                table.style = "Table Grid"
            except Exception:
                pass
            for ridx, row_cells in enumerate(parsed):
                row = table.add_row().cells
                padded = row_cells + [""] * (col_count - len(row_cells))
                for cidx, val in enumerate(padded):
                    row[cidx].text = val
                    if ridx == 0:
                        for run in row[cidx].paragraphs[0].runs:
                            run.bold = True
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(_strip_inline_md(stripped[4:]), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(_strip_inline_md(stripped[3:]), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(_strip_inline_md(stripped[2:]), level=1)
            i += 1
            continue

        # Bullets ("- " or "* ")
        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(_strip_inline_md(stripped[2:]), style="List Bullet")
            i += 1
            continue

        doc.add_paragraph(_strip_inline_md(stripped))
        i += 1

    try:
        doc.save(str(out_path))
        return out_path.exists()
    except Exception:
        return False
